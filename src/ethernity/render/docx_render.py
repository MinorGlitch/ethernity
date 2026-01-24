#!/usr/bin/env python3
# Copyright (C) 2026 Alex Stoyanov
#
# This program is free software: you can redistribute it and/or modify
# it under the terms of the GNU General Public License as published by
# the Free Software Foundation; either version 3 of the License, or
# (at your option) any later version.
#
# This program is distributed in the hope that it will be useful,
# but WITHOUT ANY WARRANTY; without even the implied warranty of
# MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE.  See the
# GNU General Public License for more details.
#
# You should have received a copy of the GNU General Public License along with this program.
# If not, see <https://www.gnu.org/licenses/>.

from __future__ import annotations

from pathlib import Path
from typing import Final

from docx import Document as create_docx_document
from docx.enum.section import WD_ORIENTATION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm

from .storage_paths import (
    DEFAULT_LOGO_PATH,
    EnvelopeKind,
    EnvelopeOrientation,
    envelope_page_size_mm,
)

_DEFAULT_LOGO_WIDTH_MM: Final[float] = 100.0


def render_envelope_docx(
    output_path: str | Path,
    *,
    kind: EnvelopeKind,
    logo_path: str | Path | None = None,
    orientation: EnvelopeOrientation = "portrait",
    border_inset_mm: float = 7.0,
    border_size_eighth_pt: int = 8,
    logo_width_mm: float = _DEFAULT_LOGO_WIDTH_MM,
) -> Path:
    """Render an envelope DOCX with a centered logo and a border."""
    output_path = Path(output_path)
    resolved_logo_path = Path(logo_path) if logo_path is not None else DEFAULT_LOGO_PATH
    if not resolved_logo_path.exists():
        raise FileNotFoundError(f"logo file not found: {resolved_logo_path}")

    doc = create_docx_document()
    _remove_leading_empty_paragraph(doc)

    page_w_mm, page_h_mm = envelope_page_size_mm(kind, orientation)

    section = doc.sections[0]
    section.orientation = (
        WD_ORIENTATION.LANDSCAPE if orientation == "landscape" else WD_ORIENTATION.PORTRAIT
    )
    section.page_width = Mm(page_w_mm)
    section.page_height = Mm(page_h_mm)

    inset = Mm(border_inset_mm)
    section.top_margin = inset
    section.bottom_margin = inset
    section.left_margin = inset
    section.right_margin = inset

    content_w_mm = max(1.0, page_w_mm - 2.0 * border_inset_mm)
    content_h_mm = max(1.0, page_h_mm - 2.0 * border_inset_mm)
    logo_width_mm = min(float(logo_width_mm), content_w_mm)

    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = Mm(content_w_mm)
    _set_table_border(table, border_size_eighth_pt=border_size_eighth_pt)

    row = table.rows[0]
    row.height = Mm(content_h_mm)
    row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

    cell = row.cells[0]
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    _set_cell_margins_zero(cell)

    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = 0
    paragraph.paragraph_format.space_after = 0

    run = paragraph.add_run()
    run.add_picture(str(resolved_logo_path), width=Mm(logo_width_mm))

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path


def render_envelope_c6_docx(
    output_path: str | Path,
    *,
    logo_path: str | Path | None = None,
    orientation: EnvelopeOrientation = "portrait",
    border_inset_mm: float = 7.0,
    border_size_eighth_pt: int = 8,
    logo_width_mm: float = _DEFAULT_LOGO_WIDTH_MM,
) -> Path:
    return render_envelope_docx(
        output_path,
        kind="c6",
        logo_path=logo_path,
        orientation=orientation,
        border_inset_mm=border_inset_mm,
        border_size_eighth_pt=border_size_eighth_pt,
        logo_width_mm=logo_width_mm,
    )


def render_envelope_c5_docx(
    output_path: str | Path,
    *,
    logo_path: str | Path | None = None,
    orientation: EnvelopeOrientation = "portrait",
    border_inset_mm: float = 7.0,
    border_size_eighth_pt: int = 8,
    logo_width_mm: float = _DEFAULT_LOGO_WIDTH_MM,
) -> Path:
    return render_envelope_docx(
        output_path,
        kind="c5",
        logo_path=logo_path,
        orientation=orientation,
        border_inset_mm=border_inset_mm,
        border_size_eighth_pt=border_size_eighth_pt,
        logo_width_mm=logo_width_mm,
    )


def render_envelope_c4_docx(
    output_path: str | Path,
    *,
    logo_path: str | Path | None = None,
    orientation: EnvelopeOrientation = "portrait",
    border_inset_mm: float = 7.0,
    border_size_eighth_pt: int = 8,
    logo_width_mm: float = _DEFAULT_LOGO_WIDTH_MM,
) -> Path:
    return render_envelope_docx(
        output_path,
        kind="c4",
        logo_path=logo_path,
        orientation=orientation,
        border_inset_mm=border_inset_mm,
        border_size_eighth_pt=border_size_eighth_pt,
        logo_width_mm=logo_width_mm,
    )


def render_envelope_dl_docx(
    output_path: str | Path,
    *,
    logo_path: str | Path | None = None,
    orientation: EnvelopeOrientation = "portrait",
    border_inset_mm: float = 7.0,
    border_size_eighth_pt: int = 8,
    logo_width_mm: float = _DEFAULT_LOGO_WIDTH_MM,
) -> Path:
    return render_envelope_docx(
        output_path,
        kind="dl",
        logo_path=logo_path,
        orientation=orientation,
        border_inset_mm=border_inset_mm,
        border_size_eighth_pt=border_size_eighth_pt,
        logo_width_mm=logo_width_mm,
    )


def _remove_leading_empty_paragraph(doc: object) -> None:
    paragraphs = getattr(doc, "paragraphs", [])
    if not paragraphs:
        return
    paragraph = paragraphs[0]
    text = getattr(paragraph, "text", "")
    if text and str(text).strip():
        return
    element = getattr(paragraph, "_element", None)
    if element is None:
        return
    parent = element.getparent()
    if parent is None:
        return
    parent.remove(element)


def _set_table_border(table: object, *, border_size_eighth_pt: int) -> None:
    tbl = getattr(table, "_tbl", None)
    if tbl is None:
        return
    tbl_pr = getattr(tbl, "tblPr", None)
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)
    for existing in tbl_pr.findall(qn("w:tblBorders")):
        tbl_pr.remove(existing)
    tbl_borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), str(int(border_size_eighth_pt)))
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")
        tbl_borders.append(element)
    for edge in ("insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "nil")
        tbl_borders.append(element)
    tbl_pr.append(tbl_borders)


def _set_cell_margins_zero(cell: object) -> None:
    tc = getattr(cell, "_tc", None)
    if tc is None:
        return
    tc_pr = tc.get_or_add_tcPr()
    for existing in tc_pr.findall(qn("w:tcMar")):
        tc_pr.remove(existing)
    tc_mar = OxmlElement("w:tcMar")
    for side in ("top", "left", "bottom", "right"):
        node = OxmlElement(f"w:{side}")
        node.set(qn("w:w"), "0")
        node.set(qn("w:type"), "dxa")
        tc_mar.append(node)
    tc_pr.append(tc_mar)


__all__ = [
    "EnvelopeOrientation",
    "render_envelope_c4_docx",
    "render_envelope_c5_docx",
    "render_envelope_c6_docx",
    "render_envelope_dl_docx",
    "render_envelope_docx",
]
