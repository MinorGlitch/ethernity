# Copyright (C) 2026 Alex Stoyanov
#
# This program is free software: you can redistribute it and/or modify
# it under the terms of the GNU General Public License as published by
# the Free Software Foundation; either version 3 of the License, or
# (at your option) any later version.
#
# This program is distributed in the hope that it will be useful,
# but WITHOUT ANY WARRANTY; without even the implied warranty of
# MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE.  See the
# GNU General Public License for more details.
#
# You should have received a copy of the GNU General Public License along with this program.
# If not, see <https://www.gnu.org/licenses/>.

from __future__ import annotations

import tempfile
import unittest
from pathlib import Path
from types import SimpleNamespace
from unittest import mock

from docx import Document
from docx.enum.section import WD_ORIENTATION
from docx.oxml.ns import qn
from docx.shared import Mm

from ethernity.render import docx_render
from ethernity.render.storage_paths import DEFAULT_LOGO_PATH


class TestDocxRender(unittest.TestCase):
    def test_render_envelope_docx_missing_logo_raises(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            output_path = Path(tmpdir) / "envelope.docx"
            missing_logo = Path(tmpdir) / "missing-logo.png"
            with self.assertRaisesRegex(FileNotFoundError, "logo file not found"):
                docx_render.render_envelope_docx(
                    output_path,
                    kind="c6",
                    logo_path=missing_logo,
                )

    def test_render_envelope_docx_success_portrait_and_landscape(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            for orientation, expected in (
                ("portrait", WD_ORIENTATION.PORTRAIT),
                ("landscape", WD_ORIENTATION.LANDSCAPE),
            ):
                with self.subTest(orientation=orientation):
                    output_path = Path(tmpdir) / f"envelope-{orientation}.docx"
                    rendered = docx_render.render_envelope_docx(
                        output_path,
                        kind="c6",
                        logo_path=DEFAULT_LOGO_PATH,
                        orientation=orientation,
                        border_inset_mm=7.0,
                        border_size_eighth_pt=10,
                        logo_width_mm=90.0,
                    )
                    self.assertEqual(rendered, output_path)
                    self.assertTrue(output_path.is_file())

                    doc = Document(output_path)
                    section = doc.sections[0]
                    self.assertEqual(section.orientation, expected)
                    self.assertAlmostEqual(int(section.top_margin), int(Mm(7.0)), delta=200)
                    self.assertAlmostEqual(int(section.left_margin), int(Mm(7.0)), delta=200)
                    self.assertEqual(len(doc.inline_shapes), 1)

    def test_render_envelope_docx_clamps_content_and_logo_width(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            output_path = Path(tmpdir) / "clamped.docx"
            docx_render.render_envelope_docx(
                output_path,
                kind="c6",
                logo_path=DEFAULT_LOGO_PATH,
                border_inset_mm=1000.0,
                logo_width_mm=9999.0,
            )

            doc = Document(output_path)
            self.assertEqual(len(doc.inline_shapes), 1)
            self.assertEqual(doc.inline_shapes[0].width, Mm(1.0))

    def test_wrapper_functions_forward_kind_and_parameters(self) -> None:
        with mock.patch.object(
            docx_render, "render_envelope_docx", return_value=Path("ok.docx")
        ) as render_mock:
            self.assertEqual(
                docx_render.render_envelope_c6_docx("a.docx", logo_path="logo.png"),
                Path("ok.docx"),
            )
            self.assertEqual(
                docx_render.render_envelope_c5_docx("b.docx", orientation="landscape"),
                Path("ok.docx"),
            )
            self.assertEqual(
                docx_render.render_envelope_c4_docx("c.docx", border_inset_mm=9.0),
                Path("ok.docx"),
            )
            self.assertEqual(
                docx_render.render_envelope_dl_docx("d.docx", logo_width_mm=55.0),
                Path("ok.docx"),
            )

        self.assertEqual(render_mock.call_args_list[0].kwargs["kind"], "c6")
        self.assertEqual(render_mock.call_args_list[1].kwargs["kind"], "c5")
        self.assertEqual(render_mock.call_args_list[2].kwargs["kind"], "c4")
        self.assertEqual(render_mock.call_args_list[3].kwargs["kind"], "dl")

    def test_remove_leading_empty_paragraph_branches(self) -> None:
        docx_render._remove_leading_empty_paragraph(SimpleNamespace(paragraphs=[]))
        docx_render._remove_leading_empty_paragraph(
            SimpleNamespace(paragraphs=[SimpleNamespace(text="filled", _element=mock.Mock())])
        )
        docx_render._remove_leading_empty_paragraph(
            SimpleNamespace(paragraphs=[SimpleNamespace(text=" ")])
        )

        no_parent_element = mock.Mock()
        no_parent_element.getparent.return_value = None
        docx_render._remove_leading_empty_paragraph(
            SimpleNamespace(paragraphs=[SimpleNamespace(text=" ", _element=no_parent_element)])
        )
        no_parent_element.getparent.assert_called_once()

        parent = mock.Mock()
        element = mock.Mock()
        element.getparent.return_value = parent
        docx_render._remove_leading_empty_paragraph(
            SimpleNamespace(paragraphs=[SimpleNamespace(text=" ", _element=element)])
        )
        parent.remove.assert_called_once_with(element)

    def test_set_table_border_branches_and_replaces_existing(self) -> None:
        docx_render._set_table_border(object(), border_size_eighth_pt=8)

        fake_tbl = SimpleNamespace(tblPr=None)
        fake_tbl.insert = mock.Mock(side_effect=lambda _idx, node: setattr(fake_tbl, "tblPr", node))
        fake_table = SimpleNamespace(_tbl=fake_tbl)
        docx_render._set_table_border(fake_table, border_size_eighth_pt=6)
        fake_tbl.insert.assert_called_once()
        self.assertIsNotNone(fake_tbl.tblPr)

        doc = Document()
        table = doc.add_table(rows=1, cols=1)
        docx_render._set_table_border(table, border_size_eighth_pt=7)
        docx_render._set_table_border(table, border_size_eighth_pt=4)

        borders = table._tbl.tblPr.findall(qn("w:tblBorders"))
        self.assertEqual(len(borders), 1)
        top = borders[0].find(qn("w:top"))
        self.assertIsNotNone(top)
        assert top is not None
        self.assertEqual(top.get(qn("w:sz")), "4")

    def test_set_cell_margins_zero_branches_and_replaces_existing(self) -> None:
        docx_render._set_cell_margins_zero(object())

        doc = Document()
        table = doc.add_table(rows=1, cols=1)
        cell = table.rows[0].cells[0]
        docx_render._set_cell_margins_zero(cell)
        docx_render._set_cell_margins_zero(cell)

        tc_mar_nodes = cell._tc.get_or_add_tcPr().findall(qn("w:tcMar"))
        self.assertEqual(len(tc_mar_nodes), 1)
        for side in ("top", "left", "bottom", "right"):
            side_node = tc_mar_nodes[0].find(qn(f"w:{side}"))
            self.assertIsNotNone(side_node)
            assert side_node is not None
            self.assertEqual(side_node.get(qn("w:w")), "0")
            self.assertEqual(side_node.get(qn("w:type")), "dxa")


if __name__ == "__main__":
    unittest.main()
